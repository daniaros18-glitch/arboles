# -*- coding: utf-8 -*-
"""Convierte un documento Markdown del proyecto a Word (.docx) legible.

Soporta: titulos (# a ####), parrafos, negrita, cursiva, codigo en linea,
enlaces, tablas, listas con viñeta y numeradas, citas y lineas divisorias.

Uso:
    python scripts/md_a_word.py entrada.md [salida.docx]
"""
import os
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

GRIS = RGBColor(0x44, 0x44, 0x44)
VERDE = RGBColor(0x1B, 0x4D, 0x3E)
ANCHO_UTIL = Inches(6.5)


def sombrear(celda, color):
    """Pinta el fondo de una celda (python-docx no lo expone directamente)."""
    tc = celda._tc.get_or_add_tcPr()
    sombra = OxmlElement('w:shd')
    sombra.set(qn('w:val'), 'clear')
    sombra.set(qn('w:fill'), color)
    tc.append(sombra)


def escribir_inline(parrafo, texto):
    """Interpreta **negrita**, *cursiva*, `codigo` y [enlace](url)."""
    texto = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'\1', texto)   # enlace -> solo el texto
    partes = re.split(r'(\*\*[^*]+\*\*|`[^`]+`|(?<!\*)\*[^*]+\*(?!\*))', texto)
    for parte in partes:
        if not parte:
            continue
        if parte.startswith('**') and parte.endswith('**'):
            parrafo.add_run(parte[2:-2]).bold = True
        elif parte.startswith('`') and parte.endswith('`'):
            r = parrafo.add_run(parte[1:-1])
            r.font.name = 'Consolas'
            r.font.size = Pt(9.5)
        elif parte.startswith('*') and parte.endswith('*'):
            parrafo.add_run(parte[1:-1]).italic = True
        else:
            parrafo.add_run(parte)


def fila_de_tabla(linea):
    celdas = [c.strip() for c in linea.strip().strip('|').split('|')]
    return celdas


def es_separador(linea):
    return bool(re.match(r'^\s*\|?[\s:\-|]+\|[\s:\-|]*$', linea)) and '-' in linea


def convertir(ruta_md, ruta_docx):
    with open(ruta_md, encoding='utf-8') as f:
        lineas = f.read().split('\n')

    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10.5)

    i = 0
    while i < len(lineas):
        linea = lineas[i]
        despojada = linea.strip()

        # --- tabla ---
        if despojada.startswith('|') and i + 1 < len(lineas) and es_separador(lineas[i + 1]):
            encabezados = fila_de_tabla(despojada)
            filas = []
            i += 2
            while i < len(lineas) and lineas[i].strip().startswith('|'):
                filas.append(fila_de_tabla(lineas[i].strip()))
                i += 1
            tabla = doc.add_table(rows=1, cols=len(encabezados))
            tabla.style = 'Table Grid'
            tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
            for j, texto in enumerate(encabezados):
                celda = tabla.rows[0].cells[j]
                celda.text = ''
                p = celda.paragraphs[0]
                escribir_inline(p, texto)
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(9.5)
                sombrear(celda, 'E8F0E4')
            for fila in filas:
                celdas = tabla.add_row().cells
                for j, texto in enumerate(fila[:len(encabezados)]):
                    celdas[j].text = ''
                    p = celdas[j].paragraphs[0]
                    escribir_inline(p, texto)
                    for r in p.runs:
                        r.font.size = Pt(9.5)
            doc.add_paragraph()
            continue

        # --- titulos ---
        m = re.match(r'^(#{1,4})\s+(.*)$', despojada)
        if m:
            nivel = len(m.group(1))
            titulo = re.sub(r'[#*`]', '', m.group(2)).strip()
            h = doc.add_heading('', level=min(nivel, 4))
            r = h.add_run(titulo)
            r.font.color.rgb = VERDE
            i += 1
            continue

        # --- divisoria ---
        if re.match(r'^\s*---+\s*$', linea):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(6)
            pr = p._p.get_or_add_pPr()
            borde = OxmlElement('w:pBdr')
            abajo = OxmlElement('w:bottom')
            abajo.set(qn('w:val'), 'single')
            abajo.set(qn('w:sz'), '6')
            abajo.set(qn('w:color'), 'BBBBBB')
            borde.append(abajo)
            pr.append(borde)
            i += 1
            continue

        # --- cita ---
        if despojada.startswith('>'):
            texto = despojada.lstrip('>').strip()
            if texto:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                escribir_inline(p, texto)
                for r in p.runs:
                    r.italic = True
                    r.font.color.rgb = GRIS
            i += 1
            continue

        # --- listas ---
        m = re.match(r'^(\s*)[-*]\s+(.*)$', linea)
        if m:
            p = doc.add_paragraph(style='List Bullet')
            if len(m.group(1)) >= 2:
                p.paragraph_format.left_indent = Inches(0.65)
            escribir_inline(p, m.group(2))
            i += 1
            continue
        m = re.match(r'^(\s*)\d+\.\s+(.*)$', linea)
        if m:
            p = doc.add_paragraph(style='List Number')
            escribir_inline(p, m.group(2))
            i += 1
            continue

        # --- parrafo o linea en blanco ---
        if despojada:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            escribir_inline(p, despojada)
        i += 1

    doc.save(ruta_docx)
    return ruta_docx


if __name__ == '__main__':
    if len(sys.argv) < 2:
        raise SystemExit('Uso: python scripts/md_a_word.py entrada.md [salida.docx]')
    entrada = sys.argv[1]
    salida = sys.argv[2] if len(sys.argv) > 2 else os.path.splitext(entrada)[0] + '.docx'
    print('Creado:', convertir(entrada, salida))
